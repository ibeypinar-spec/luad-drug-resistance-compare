"""
Generate LUAD_manuscript_v1.1.docx from manuscript_draft_v1.md
Journal submission format: text → references → tables → figures
"""
import sys, os, re

if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx ...")
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

WORK_DIR = "G:\\Drive'\u0131m\\\u00c7al\u0131\u015fma\\Aktif \u00c7al\u0131\u015fmalar\\9. A\u00e7\u0131k eri\u015fim veri \u00e7al\u0131\u015fmalar\u0131\\LUAD Drug Resistance Study"
RESULTS_DIR = os.path.join(WORK_DIR, "results")
OUTPUT = os.path.join(WORK_DIR, "LUAD_manuscript_v1.1.docx")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_margins(doc):
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

def set_normal_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(24)  # double-spaced

def add_runs(para, text):
    """Parse **bold**, *italic*, and ***bold-italic*** inline markers."""
    parts = re.split(r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = para.add_run(part[3:-3])
            run.bold = True; run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)

def h(doc, text, level):
    p = doc.add_heading('', level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    return p

def para(doc, text, center=False, bold=False, italic=False, indent=True):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    add_runs(p, text)
    if bold:
        for run in p.runs: run.bold = True
    if italic:
        for run in p.runs: run.italic = True
    return p

def label_para(doc, label, text):
    """Paragraph with a bold label like 'Background: ' followed by normal text."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(label)
    r.bold = True
    add_runs(p, text)
    return p

def page_break(doc):
    doc.add_page_break()

def add_image_safe(doc, path, width_inches=6.0):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Figure file not found: {os.path.basename(path)}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs: run.italic = True

def hr(doc):
    """Thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
doc = Document()
set_margins(doc)
set_normal_style(doc)

# ── TITLE PAGE ──────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(24)
r = tp.add_run("pARACNe-Inferred Tyrosine Kinase Network Predicts Drug Resistance\nin Lung Adenocarcinoma: A Multi-Dataset Integration Study")
r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'

doc.add_paragraph()
rp = doc.add_paragraph()
rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_runs(rp, "*Running title:* Kinase Network and Drug Resistance in LUAD")

doc.add_paragraph()
ap = doc.add_paragraph()
ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
ap.paragraph_format.first_line_indent = Cm(0)
r = ap.add_run("[Author names]")
r.font.name = 'Times New Roman'

affp = doc.add_paragraph()
affp.alignment = WD_ALIGN_PARAGRAPH.CENTER
affp.paragraph_format.first_line_indent = Cm(0)
r = affp.add_run("Alanya Alaaddin Keykubat University Faculty of Medicine, "
                  "Department of Medical Oncology, Alanya, Turkey")
r.font.name = 'Times New Roman'

corrp = doc.add_paragraph()
corrp.alignment = WD_ALIGN_PARAGRAPH.CENTER
corrp.paragraph_format.first_line_indent = Cm(0)
add_runs(corrp, "*Corresponding author:* Dr. İso — [email]")

doc.add_paragraph()
jp = doc.add_paragraph()
jp.alignment = WD_ALIGN_PARAGRAPH.CENTER
jp.paragraph_format.first_line_indent = Cm(0)
add_runs(jp, "*Target journal:* Journal of Thoracic Oncology (IF ~22)")

kp = doc.add_paragraph()
kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
kp.paragraph_format.first_line_indent = Cm(0)
add_runs(kp, "**Keywords:** lung adenocarcinoma, tyrosine kinase, drug resistance, "
             "pARACNe, network pharmacology, YES1, FAK, proteogenomics")

page_break(doc)

# ── ABSTRACT ────────────────────────────────────────────────────────────────
h(doc, "STRUCTURED ABSTRACT", 1)

label_para(doc, "Background: ",
    "Drug resistance remains a critical barrier in lung adenocarcinoma (LUAD), particularly "
    "to EGFR tyrosine kinase inhibitors (TKIs). Proteomic network analysis via pARACNe has "
    "defined kinase\u2013substrate interaction maps in LUAD, yet the pharmacological relevance "
    "of these networks remains unexplored at scale.")

label_para(doc, "Methods: ",
    "We integrated four independent datasets: (1) the pARACNe LUAD proteome-signaling network "
    "(46 tyrosine kinases, 2,064 kinase\u2013substrate interactions); (2) GDSC v2 drug sensitivity "
    "profiles (62 LUAD cell lines, 286 compounds); (3) TCGA-LUAD clinical and transcriptomic data "
    "(n\u00a0=\u00a0497 patients); and (4) CPTAC-LUAD proteogenomic data (n\u00a0=\u00a0110 patients, "
    "10,316 proteins, 40,893 phosphosites). For each hub kinase, we computed expression\u2013IC50 "
    "Spearman correlations, Kaplan\u2013Meier overall survival analysis, multivariate Cox regression, "
    "and a substrate-based kinase activation score (KAS) from phosphoproteomic data.")

label_para(doc, "Results: ",
    "PTK2/FAK emerged as the highest-connectivity hub (195 substrates), followed by MET (168), "
    "LCK (138), YES1 (136), FYN (134), and EGFR (123). In TCGA-LUAD, YES1 high expression was "
    "independently associated with poor overall survival in multivariate Cox regression "
    "(HR\u00a0=\u00a01.28, 95% CI 1.03\u20131.60, p\u00a0=\u00a00.028). Tertile-based Kaplan\u2013Meier "
    "analysis additionally identified LCK (FDR\u00a0=\u00a00.018) and YES1 (FDR\u00a0=\u00a00.027) as "
    "significantly associated with OS. Full GDSC2 screening (2,574 tests) identified nominal "
    "kinase\u2013drug associations including SRC\u2013bosutinib (\u03c1\u00a0=\u00a0\u22120.36, "
    "p\u00a0=\u00a00.009) and EGFR\u2013gemcitabine (\u03c1\u00a0=\u00a00.41, p\u00a0=\u00a00.003); no "
    "association survived FDR correction. CPTAC validation confirmed 37/752 (4.9%) pARACNe-predicted "
    "kinase\u2013substrate pairs (FDR\u00a0<\u00a00.05). SRC-family kinases exhibited near-perfect "
    "co-activation (YES1\u2013FYN: \u03c1\u00a0=\u00a00.999; LYN\u2013YES1: \u03c1\u00a0=\u00a00.912), "
    "with PC1 of the nine-kinase KAS matrix explaining 80.8% of variance.")

label_para(doc, "Conclusions: ",
    "Multi-dataset integration of the pARACNe LUAD signaling network identifies YES1 as an "
    "independent prognostic biomarker and reveals SRC-family co-activation as a potential mechanism "
    "of TKI resistance. These findings provide a network-level rationale for combined EGFR + "
    "SRC-family inhibition strategies in LUAD.")

page_break(doc)

# ── INTRODUCTION ────────────────────────────────────────────────────────────
h(doc, "1. INTRODUCTION", 1)

para(doc,
    "Lung adenocarcinoma (LUAD) accounts for approximately 40% of all lung cancers and is the "
    "leading cause of cancer-related mortality worldwide [1]. The identification of druggable driver "
    "mutations\u2014most notably EGFR activating mutations (present in 15\u201340% of LUAD depending "
    "on ethnic background) and ALK/ROS1 fusions\u2014has transformed treatment paradigms, with EGFR "
    "TKIs achieving objective response rates of 60\u201370% [2]. However, resistance invariably "
    "develops, with median progression-free survival of 10\u201318 months on first- and "
    "second-generation EGFR TKIs and up to 18 months on third-generation osimertinib [3]. The T790M "
    "gatekeeper mutation accounts for approximately 50% of acquired resistance to first-generation "
    "TKIs, while MET amplification, HER2 amplification, and bypass signaling through the SRC/FAK "
    "axis account for a substantial proportion of the remainder [4].")

para(doc,
    "The complexity of resistance mechanisms reflects the extensive crosstalk within receptor "
    "tyrosine kinase (RTK) signaling networks. Once EGFR is inhibited, cancer cells can rewire "
    "signaling through alternative kinase hubs to maintain proliferative and survival signals\u2014"
    "a phenomenon termed bypass resistance [CITE]. Despite this understanding, the systematic "
    "mapping of which kinase nodes drive resistance across the LUAD signaling landscape has remained "
    "incomplete, owing to the difficulty of capturing network-level interactions from gene-by-gene "
    "studies.")

para(doc,
    "Computational network inference has emerged as a powerful approach to reconstruct "
    "kinase\u2013substrate signaling from proteomic data. The pARACNe algorithm, applied to the "
    "Columbia LUAD proteome dataset, identified 46 tyrosine kinases and 2,064 high-confidence "
    "kinase\u2013substrate interactions specific to LUAD tissue [5]. This network provides a unique "
    "resource to prioritize kinase hubs based on network connectivity\u2014an orthogonal approach "
    "to traditional mutation-centric analyses. However, the pharmacological relevance of "
    "pARACNe-predicted hub kinases has not been systematically evaluated against drug sensitivity "
    "data or clinical outcomes.")

para(doc,
    "Here, we present the first large-scale integration of the pARACNe LUAD signaling network with: "
    "(i) GDSC v2 drug sensitivity profiles across 62 LUAD cell lines and 286 compounds; "
    "(ii) clinical outcome data from 497 TCGA-LUAD patients; and (iii) CPTAC-LUAD phosphoproteomic "
    "validation across 110 patients. Our analyses reveal YES1/SRC family kinases as network hubs "
    "with independent prognostic significance, and provide a multi-dataset framework for identifying "
    "combination therapy targets in drug-resistant LUAD.")

page_break(doc)

# ── METHODS ──────────────────────────────────────────────────────────────────
h(doc, "2. METHODS", 1)

h(doc, "2.1 Data Sources", 2)

label_para(doc, "pARACNe LUAD Signaling Network. ",
    "The LUAD proteome-signaling network was obtained from the Columbia University LUAD study "
    "(Bansal et al., 2019, PLoS ONE) [5]. The protein-level network (S2) contains 46 tyrosine "
    "kinase regulators and 2,064 kinase\u2013substrate interactions, with mutual information (MI) "
    "scores reflecting interaction confidence. All analyses used the S2 protein-level network.")

label_para(doc, "GDSC v2 Drug Sensitivity. ",
    "Dose\u2013response data were obtained from the Genomics of Drug Sensitivity in Cancer database "
    "version 2 (GDSC2, release 8.5; www.cancerrxgene.org) [6]. The dataset provides fitted "
    "dose-response parameters, including LN_IC50 (natural log of IC50 in \u03bcM) and AUC, for "
    "62 LUAD cell lines and 286 compounds. Cell line identity was cross-referenced to DepMap "
    "using SANGER_MODEL_ID identifiers.")

label_para(doc, "DepMap Gene Expression. ",
    "RNA-seq expression data (log2(TPM+1)) for DepMap 23Q4 were downloaded from the Broad "
    "Institute DepMap portal (figshare DOI: 10.6084/m9.figshare.24667905) [7]. LUAD cell lines "
    "were identified using OncotreeCode\u00a0=\u00a0'LUAD' (n\u00a0=\u00a089 cell lines), of which "
    "62 had matching GDSC2 drug sensitivity data.")

label_para(doc, "TCGA-LUAD. ",
    "Clinical and mRNA expression data (RNA-Seq V2 RSEM) for the TCGA-LUAD PanCancer Atlas "
    "cohort (n\u00a0=\u00a0511 patients with expression; n\u00a0=\u00a0566 with clinical data) were "
    "downloaded from cBioPortal (study ID: luad_tcga_pan_can_atlas_2018) [9]. For survival "
    "analyses, patients with available mRNA expression and overall survival data were included "
    "(n\u00a0=\u00a0497).")

label_para(doc, "CPTAC-LUAD. ",
    "Processed proteome and phosphoproteome ratio data (log2 tumor/reference ratio, NA-removed) "
    "and clinical annotations for 110 LUAD tumor specimens were downloaded from LinkedOmics "
    "(linkedomics.org; CPTAC LUAD Discovery Study) [8]. The proteome contained 10,316 proteins "
    "and the phosphoproteome 40,893 phosphosites across 110 patients.")

h(doc, "2.2 Hub Kinase Identification", 2)
para(doc,
    "Kinase hub scores were calculated as the number of unique substrate proteins per kinase in "
    "the pARACNe S2 protein-level network. Top hub kinases were defined as those with \u2265\u00a0100 "
    "substrate connections, yielding nine focal kinases: PTK2 (FAK), MET, LCK, YES1, FYN, EGFR, "
    "LYN, DDR1, and SRC.")

h(doc, "2.3 Kinase Expression\u2013Drug Sensitivity Correlation", 2)
para(doc,
    "For each focal kinase across 62 matched LUAD cell lines, Spearman rank correlation (\u03c1) "
    "was computed between kinase mRNA expression (log2 TPM+1) and drug LN_IC50 for all 286 "
    "GDSC2 compounds. A minimum of eight matched cell lines was required per test. Multiple "
    "testing correction was applied using the Benjamini-Hochberg false discovery rate (FDR) "
    "procedure. Negative \u03c1 indicates that higher kinase expression associates with lower IC50 "
    "(greater drug sensitivity); positive \u03c1 indicates resistance.")

h(doc, "2.4 Survival Analysis", 2)
para(doc,
    "Overall survival (OS) was used as the primary endpoint in TCGA-LUAD. Expression-based "
    "patient stratification was performed by median dichotomization and, for key kinases, tertile "
    "split (bottom vs. top 33%). Kaplan\u2013Meier survival curves were constructed, and group "
    "differences were assessed by log-rank test. For Cox proportional hazards regression, kinase "
    "expression values were included as continuous variables along with age as a covariate. All "
    "Cox models were fitted with a ridge penalty (\u03bb\u00a0=\u00a00.1) to handle collinearity "
    "among co-expressed kinases.")
para(doc,
    "EGFR mutant patients were identified from TCGA somatic mutation data (MAF format, "
    "data_mutations.txt) by filtering for non-synonymous variants in EGFR. Subgroup survival "
    "analyses were performed in the EGFR mutant cohort.")

h(doc, "2.5 Kinase Activation Score (KAS)", 2)
para(doc,
    "For each focal kinase in CPTAC-LUAD, a substrate-based Kinase Activation Score was computed "
    "as the mean log2 phosphorylation ratio across all phosphosites mapping to the kinase\u2019s "
    "pARACNe substrates. Phosphosite-to-substrate mapping was performed by extracting the gene "
    "name from each phosphosite identifier (format: GENE:RefSeq:site). A minimum of five substrate "
    "phosphosites was required per kinase. A composite multi-kinase activation score was derived by "
    "principal component analysis (PCA) of the nine individual KAS vectors; PC1 was used as the "
    "composite score.")

h(doc, "2.6 pARACNe Network Validation", 2)
para(doc,
    "Kinase\u2013substrate interactions predicted by pARACNe were validated in CPTAC-LUAD by "
    "testing whether kinase protein abundance (proteome) correlated with substrate phosphorylation "
    "level (phosphoproteome) across 110 patients. Spearman correlation was computed for each "
    "kinase\u2013substrate\u2013site triplet (up to 15 substrates \u00d7 all matching phosphosites "
    "per substrate per kinase). FDR correction was applied across all tests.")

h(doc, "2.7 Statistical Analysis", 2)
para(doc,
    "All statistical analyses were performed in Python 3.13 using pandas, scipy, statsmodels, "
    "scikit-learn, and lifelines. Figures were generated with matplotlib and seaborn. FDR "
    "correction used the Benjamini-Hochberg method throughout. Statistical significance was defined "
    "as FDR\u00a0<\u00a00.05; borderline associations were noted at p\u00a0<\u00a00.1 (uncorrected). "
    "All analyses are fully reproducible; code is available at [GitHub URL].")

page_break(doc)

# ── RESULTS ───────────────────────────────────────────────────────────────────
h(doc, "3. RESULTS", 1)

h(doc, "3.1 pARACNe Hub Kinase Architecture in LUAD", 2)
para(doc,
    "Analysis of the pARACNe S2 protein-level network (46 tyrosine kinases, 2,064 interactions) "
    "revealed a highly skewed degree distribution, with nine kinases exceeding 100 substrate "
    "connections (Figure\u00a01). PTK2 (FAK) displayed the highest connectivity (195 substrates), "
    "followed by MET (168), LCK (138), YES1 (136), FYN (134), EGFR (123), LYN (117), DDR1 (113), "
    "and SRC (112). Notably, four of the top nine hub kinases\u2014LCK, YES1, FYN, and LYN\u2014"
    "belong to the SRC family, suggesting coordinated SRC-family activity as a defining feature of "
    "LUAD signaling architecture. All nine focal kinases were confirmed to be expressed in both "
    "GDSC2 LUAD cell lines and CPTAC tumor specimens.")

h(doc, "3.2 Kinase Expression Associates with Drug Sensitivity in LUAD Cell Lines", 2)
para(doc,
    "To assess the pharmacological relevance of hub kinases, we correlated mRNA expression of the "
    "nine focal kinases with LN_IC50 values for all 286 GDSC2 compounds across 62 LUAD cell lines "
    "(n\u00a0=\u00a044\u201353 per test after cell line matching; total 2,574 tests) "
    "(Figure\u00a02; Supplementary Table\u00a0S1).")
para(doc,
    "Across 2,574 kinase\u2013drug tests, no association survived FDR correction "
    "(minimum FDR\u00a0=\u00a00.41), consistent with limited statistical power at "
    "n\u00a0=\u00a044\u201353 cell lines per compound. Nevertheless, several nominal associations "
    "supported the biological coherence of our framework. SRC expression correlated negatively "
    "with bosutinib IC50 (\u03c1\u00a0=\u00a0\u22120.36, p\u00a0=\u00a00.009), consistent with direct "
    "SRC target engagement by bosutinib and serving as an internal positive control for the pipeline. "
    "EGFR expression positively correlated with gemcitabine IC50 (\u03c1\u00a0=\u00a00.41, "
    "p\u00a0=\u00a00.003) and oxaliplatin IC50 (\u03c1\u00a0=\u00a00.40, p\u00a0=\u00a00.003), "
    "suggesting that EGFR-high cell lines may harbor relative resistance to these chemotherapeutic "
    "agents. DDR1 expression showed a positive association with 5-fluorouracil IC50 "
    "(\u03c1\u00a0=\u00a00.36, p\u00a0=\u00a00.008) and a negative association with gefitinib IC50 "
    "(\u03c1\u00a0=\u00a0\u22120.32, p\u00a0=\u00a00.019) (Supplementary Table\u00a0S1).")
para(doc,
    "GDSC2 ANOVA analysis confirmed that EGFR mutation status was the dominant pharmacogenomic "
    "feature for osimertinib response (effect size\u00a0=\u00a03.06, FDR\u00a0<\u00a00.001; "
    "Supplementary Table\u00a0S2), validating the pharmacogenomic sensitivity of our dataset to "
    "known EGFR-targeted therapy predictors.")

h(doc, "3.3 YES1 is an Independent Prognostic Factor in TCGA-LUAD", 2)
para(doc,
    "We next evaluated whether hub kinase expression predicted clinical outcomes in the TCGA-LUAD "
    "cohort (n\u00a0=\u00a0497 patients with matched expression and survival data). Median overall "
    "survival was [X] months; 36% of patients experienced a death event during follow-up.")
para(doc,
    "Kaplan\u2013Meier analysis stratified by median expression showed that LCK (log-rank "
    "p\u00a0=\u00a00.062) and FYN (p\u00a0=\u00a00.088) approached but did not cross the significance "
    "threshold after FDR correction. Switching to tertile-based stratification (top vs. bottom 33%) "
    "to increase sensitivity for non-linear expression effects, two kinases achieved statistical "
    "significance: **LCK** (p\u00a0=\u00a00.002, FDR\u00a0=\u00a00.018) and **YES1** "
    "(p\u00a0=\u00a00.006, FDR\u00a0=\u00a00.027), with FYN showing a borderline association "
    "(p\u00a0=\u00a00.049) (Figure\u00a03A). Patients in the top YES1 expression tertile had markedly "
    "inferior OS compared to the bottom tertile, consistent with YES1 acting as a dominantly adverse "
    "prognostic factor at high expression levels. In multivariate Cox regression including all nine "
    "kinases simultaneously with age as a covariate (n\u00a0=\u00a0487), YES1 was the only kinase to "
    "achieve statistical significance (HR\u00a0=\u00a01.28, 95%\u00a0CI 1.03\u20131.60, "
    "p\u00a0=\u00a00.028; Figure\u00a03B; Table\u00a01), indicating that higher YES1 expression "
    "independently predicts shorter overall survival irrespective of co-expressed kinases.")
para(doc,
    "In the EGFR-mutant subgroup (n\u00a0=\u00a069 patients with matched expression and survival "
    "data), YES1 expression showed a consistent adverse prognostic trend, though the restricted "
    "sample size limited statistical power for definitive subgroup conclusions "
    "(Supplementary Figure\u00a0S2).")

h(doc, "3.4 SRC-Family Co-Activation Characterizes High-Risk LUAD", 2)
para(doc,
    "PCA of the nine-kinase KAS matrix in CPTAC (n\u00a0=\u00a0110 patients) revealed that PC1 "
    "explained **80.8%** of total variance (PC2: 9.2%), with all nine kinases loading positively "
    "on PC1 (Figure\u00a04A\u2013C). This dominant co-activation axis\u2014which we term the "
    "composite Kinase Activation Score (cKAS)\u2014reflects a pan-kinase activation program in "
    "which SRC-family members, EGFR, MET, and PTK2 co-vary coherently across patients, suggesting "
    "shared upstream regulatory mechanisms.")
para(doc,
    "Kinase co-activation analysis demonstrated remarkably strong within-family correlations among "
    "SRC-family members in CPTAC phosphoproteomics: YES1\u2013FYN \u03c1\u00a0=\u00a00.999, "
    "LYN\u2013YES1 \u03c1\u00a0=\u00a00.912, FYN\u2013LYN \u03c1\u00a0=\u00a00.910 (Figure\u00a05A). "
    "These near-perfect co-activation correlations indicate that YES1, FYN, and LYN are functionally "
    "coupled in LUAD tumors, likely reflecting a shared upstream activation mechanism rather than "
    "independent regulation. These co-activation patterns were recapitulated at the mRNA level in "
    "TCGA-LUAD (all SRC-family pairs \u03c1\u00a0>\u00a00.4; p\u00a0<\u00a00.001; Figure\u00a05B). "
    "The concordance across independent phosphoproteomic and transcriptomic platforms validates "
    "SRC-family co-regulation as a robust biological signal.")
para(doc,
    "KAS distributions across EGFR mutation strata in CPTAC were visualized (Figure\u00a04D), "
    "revealing heterogeneous activation levels consistent with the molecular diversity of the "
    "LUAD patient population.")

h(doc, "3.5 CPTAC Phosphoproteomics Validates pARACNe Kinase\u2013Substrate Predictions", 2)
para(doc,
    "To assess the functional validity of pARACNe predictions, we tested 752 "
    "kinase\u2013substrate\u2013site triplets in CPTAC-LUAD by correlating kinase protein abundance "
    "with substrate phosphorylation. Thirty-seven interactions (4.9%) achieved FDR\u00a0<\u00a00.05 "
    "(Figure\u00a05; Supplementary Table\u00a0S3). The most strongly validated interaction was "
    "EGFR\u00a0\u2192\u00a0DBNL:S232 (\u03c1\u00a0=\u00a00.45, FDR\u00a0=\u00a00.0006), followed by "
    "EGFR\u00a0\u2192\u00a0TJAP1:S300 (\u03c1\u00a0=\u00a00.41, FDR\u00a0=\u00a00.002) and "
    "YES1\u00a0\u2192\u00a0COPA:S173 (\u03c1\u00a0=\u00a00.37, FDR\u00a0=\u00a00.009). SRC family "
    "kinases (SRC, YES1, LYN) collectively accounted for 18 of the 37 validated interactions (49%), "
    "consistent with their elevated hub connectivity.")
para(doc,
    "The 4.9% validation rate significantly exceeds the 1% expected under a null hypothesis of "
    "random association (binomial test, p\u00a0<\u00a00.001), confirming the biological specificity "
    "of pARACNe predictions in independent LUAD patient tissue.")

page_break(doc)

# ── DISCUSSION ────────────────────────────────────────────────────────────────
h(doc, "4. DISCUSSION", 1)
para(doc,
    "We present the first systematic integration of the pARACNe LUAD proteomic signaling network "
    "with pharmacogenomic, transcriptomic, and phosphoproteomic datasets spanning clinical cohorts "
    "and preclinical models. Our principal findings are: (1) SRC-family kinases\u2014particularly "
    "YES1\u2014emerge as dominant hub kinases with independent prognostic significance; "
    "(2) pARACNe kinase\u2013substrate predictions are validated at clinically meaningful rates in "
    "independent patient tissue; and (3) network-level co-activation of SRC-family members is a "
    "coherent, reproducible signal across multiple platforms.")

label_para(doc, "YES1 as a drug resistance hub. ",
    "YES1, a member of the SRC family of non-receptor tyrosine kinases, has previously been "
    "implicated in osimertinib resistance in EGFR-mutant NSCLC through SRC-family reactivation "
    "downstream of EGFR inhibition [10]. Our findings extend this observation to the genomically "
    "heterogeneous LUAD population in two complementary ways: first, tertile-based "
    "Kaplan\u2013Meier analysis in TCGA-LUAD showed YES1 high expression associates with inferior "
    "OS (FDR\u00a0=\u00a00.027), revealing a non-linear prognostic effect most pronounced in the "
    "top expression tertile; second, multivariate Cox regression confirmed that YES1\u2019s "
    "prognostic effect is independent of the other eight co-expressed hub kinases "
    "(HR\u00a0=\u00a01.28, 95%\u00a0CI 1.03\u20131.60, p\u00a0=\u00a00.028). Taken together, these "
    "data suggest that YES1 activation may confer a general survival disadvantage in LUAD beyond "
    "the EGFR-mutant subgroup. Furthermore, YES1 ranks fourth by network connectivity in pARACNe "
    "(136 substrates), indicating that its clinical significance is matched by its centrality in "
    "the LUAD kinome.")

label_para(doc, "PTK2/FAK as a structural hub. ",
    "Despite having the highest substrate connectivity (195 substrates), PTK2 did not achieve "
    "prognostic significance in our TCGA analysis. This apparent paradox may reflect the known "
    "biological complexity of FAK signaling: FAK is active across diverse LUAD molecular subtypes, "
    "including those driven by both EGFR and KRAS, and its downstream effects may be "
    "context-dependent [11]. Alternatively, the mRNA level of PTK2 may not faithfully capture "
    "FAK kinase activity\u2014a limitation that protein-level or phosphorylation-based measurements "
    "(as captured by KAS) may overcome. Notably, PTK2 KAS was robustly computable in CPTAC "
    "(1,626 substrate phosphosites; n\u00a0=\u00a0110), providing a protein-activity-based "
    "complementary metric for future clinical association studies.")

label_para(doc, "SRC-family co-activation and bypass resistance. ",
    "The near-perfect co-activation among YES1, FYN, and LYN in CPTAC phosphoproteomics "
    "(YES1\u2013FYN \u03c1\u00a0=\u00a00.999; LYN\u2013YES1 \u03c1\u00a0=\u00a00.912; "
    "FYN\u2013LYN \u03c1\u00a0=\u00a00.910)\u2014and their replication at the mRNA level in "
    "TCGA\u2014supports the hypothesis that SRC-family members function as a tightly coordinated "
    "signaling module in LUAD. A \u03c1 approaching 1.0 between YES1 and FYN phosphorylation "
    "activities across 110 independent patient specimens is unusually strong for biological data "
    "and implies either a shared regulatory input (e.g., upstream receptor activation or scaffold "
    "complex co-localization) or direct transactivation. This has direct therapeutic implications: "
    "concurrent activation of multiple SRC-family kinases would be expected to limit the efficacy "
    "of single-agent dasatinib or other pan-SRC inhibitors unless all family members are "
    "co-targeted. Identification of upstream activators of the SRC co-activation cluster "
    "(e.g., integrin signaling, HGF/MET, EGFR-independent JAK pathways) represents a priority "
    "for future mechanistic study.")

label_para(doc, "Drug sensitivity correlations. ",
    "The GDSC2 analysis provided hypothesis-generating associations between kinase expression and "
    "drug response across 62 LUAD cell lines and 286 compounds. The absence of FDR-corrected "
    "significant associations across 2,574 tests reflects the inherent statistical constraint of "
    "44\u201353 cell lines per compound\u2014a limitation intrinsic to the GDSC2 LUAD cohort size "
    "rather than an absence of true pharmacogenomic relationships. The negative correlation of SRC "
    "expression with bosutinib IC50 (\u03c1\u00a0=\u00a0\u22120.36) confirms analytical validity, "
    "as bosutinib directly inhibits SRC. The positive correlations of EGFR expression with "
    "gemcitabine (\u03c1\u00a0=\u00a00.41) and oxaliplatin (\u03c1\u00a0=\u00a00.40) IC50 values "
    "suggest that EGFR signaling may promote chemotherapy resistance through upregulation of "
    "anti-apoptotic or survival pathways [CITE]. The association between DDR1 expression and "
    "gefitinib resistance, while requiring independent validation, is biologically plausible given "
    "DDR1\u2019s role in promoting epithelial-mesenchymal transition and EGFR independence through "
    "collagen-driven signaling [CITE].")

label_para(doc, "Limitations. ",
    "Several limitations should be acknowledged. First, the GDSC2 drug sensitivity correlations "
    "are based on a modest number of LUAD cell lines (n\u00a0=\u00a044\u201353 per drug), limiting "
    "statistical power for some analyses. Second, the TCGA-LUAD cohort contains patients treated "
    "across multiple eras without standardized EGFR-TKI exposure data, meaning that OS reflects a "
    "heterogeneous treatment background. Third, the pARACNe network was inferred from the Columbia "
    "LUAD proteome dataset (n\u00a0=\u00a046 kinases), which, while LUAD-specific, may not capture "
    "all biologically active kinase interactions. Fourth, the KAS metric assumes that substrate "
    "phosphorylation reflects kinase activity, which may be confounded by phosphatase activity or "
    "kinase-independent phosphorylation events.")

label_para(doc, "Clinical implications. ",
    "Our data provide a multi-dataset rationale for (1) monitoring YES1 expression as a "
    "complementary prognostic biomarker in LUAD, particularly in the context of EGFR TKI therapy; "
    "(2) exploring combined EGFR + SRC-family inhibition (e.g., osimertinib + dasatinib or "
    "bosutinib) in patients with high SRC-family co-activation scores; and (3) using PTK2/FAK KAS "
    "as a functional readout of the EGFR bypass resistance state in prospective correlative "
    "biomarker studies.")

page_break(doc)

# ── CONCLUSIONS ───────────────────────────────────────────────────────────────
h(doc, "5. CONCLUSIONS", 1)
para(doc,
    "Integration of the pARACNe LUAD signaling network with pharmacogenomic, clinical, and "
    "phosphoproteomic datasets identifies YES1 as an independent prognostic biomarker and reveals "
    "coordinated SRC-family kinase activation as a potential mechanism of drug resistance in LUAD. "
    "This multi-dataset approach validates pARACNe network predictions in independent patient "
    "cohorts and provides a network-pharmacology framework for rationally designing combination "
    "TKI strategies.")

page_break(doc)

# ── ACKNOWLEDGEMENTS / FUNDING / COI ─────────────────────────────────────────
h(doc, "ACKNOWLEDGEMENTS", 1)
para(doc, "[To be completed]", indent=False)

h(doc, "FUNDING", 1)
para(doc, "[To be completed]", indent=False)

h(doc, "CONFLICTS OF INTEREST", 1)
para(doc, "[To be completed]", indent=False)

h(doc, "DATA AVAILABILITY", 1)
para(doc,
    "All datasets used in this study are publicly available. pARACNe LUAD network: Bansal 2019 "
    "(DOI: 10.1371/journal.pone.0208646). GDSC v2: www.cancerrxgene.org. "
    "DepMap 23Q4: https://doi.org/10.6084/m9.figshare.24667905. "
    "TCGA-LUAD: cBioPortal study ID luad_tcga_pan_can_atlas_2018. "
    "CPTAC-LUAD: LinkedOmics (linkedomics.org). Analysis code: [GitHub URL \u2014 to be created].",
    indent=False)

page_break(doc)

# ── REFERENCES ────────────────────────────────────────────────────────────────
h(doc, "REFERENCES", 1)

refs = [
    "Sung H, Ferlay J, Siegel RL, et al. Global cancer statistics 2020: GLOBOCAN estimates of "
    "incidence and mortality worldwide for 36 cancers in 185 countries. CA Cancer J Clin. "
    "2021;71(3):209\u2013249.",

    "Mok TS, Wu YL, Thongprasert S, et al. Gefitinib or carboplatin\u2013paclitaxel in pulmonary "
    "adenocarcinoma. N Engl J Med. 2009;361(10):947\u2013957.",

    "Ramalingam SS, Vansteenkiste J, Planchard D, et al. Overall survival with osimertinib in "
    "untreated, EGFR-mutated advanced NSCLC. N Engl J Med. 2020;382(1):41\u201350.",

    "Oxnard GR, Hu Y, Mileham KF, et al. Assessment of resistance mechanisms and clinical "
    "implications in patients with EGFR T790M-positive lung cancer and acquired resistance to "
    "osimertinib. JAMA Oncol. 2018;4(11):1527\u20131534.",

    "Bansal M, Kinkade R, Bhardwaj A, et al. Elucidating synergistic dependencies in lung "
    "adenocarcinoma by proteome-wide signaling-network analysis. PLoS One. "
    "2019;14(1):e0208646.",

    "Iorio F, Knijnenburg TA, Vis DJ, et al. A landscape of pharmacogenomic interactions in "
    "cancer. Cell. 2016;166(3):740\u2013754.",

    "Ghandi M, Huang FW, Jan\u00e9-Valbuena J, et al. Next-generation characterization of the "
    "Cancer Cell Line Encyclopedia. Nature. 2019;569(7757):503\u2013508.",

    "Gillette MA, Satpathy S, Cao S, et al. Proteogenomic characterization reveals therapeutic "
    "vulnerabilities in lung adenocarcinoma. Cell. 2020;182(1):200\u2013225.",

    "Cancer Genome Atlas Research Network. Comprehensive molecular profiling of lung "
    "adenocarcinoma. Nature. 2014;511(7511):543\u2013550.",

    "Koga T, Kobayashi Y, Tomizawa K, et al. YES1-activated non-small cell lung cancer is "
    "responsive to dasatinib. Clin Cancer Res. 2020;26(20):5364\u20135374.",

    "Sulzmaier FJ, Jean C, Schlaepfer DD. FAK in cancer: mechanistic findings and clinical "
    "applications. Nat Rev Cancer. 2014;14(9):598\u2013610.",

    "Provencio M, Coves-Sarto J, Cou\u00f1ago F, et al. Lung adenocarcinoma with EGFR mutation: "
    "current treatment options. Cancers. 2022;14(7):1762.",
]

for i, ref in enumerate(refs, 1):
    rp = doc.add_paragraph(style='Normal')
    rp.paragraph_format.first_line_indent = Cm(-0.75)
    rp.paragraph_format.left_indent = Cm(0.75)
    run_num = rp.add_run(f"{i}.\u00a0")
    run_num.bold = True
    rp.add_run(ref)

page_break(doc)

# ── TABLES ────────────────────────────────────────────────────────────────────
h(doc, "TABLES", 1)

# Table 1
h(doc, "Table 1. Multivariate Cox Proportional Hazards Regression \u2014 TCGA-LUAD (n\u00a0=\u00a0487)", 2)
para(doc,
    "Cox PH regression with ridge penalty (\u03bb\u00a0=\u00a00.1). "
    "All nine focal kinases and age included simultaneously. "
    "Bold: statistically significant (p\u00a0<\u00a00.05).",
    indent=False)

t1_data = [
    ("Variable", "Hazard Ratio", "95% CI", "p-value"),
    ("PTK2", "0.82", "0.63\u20131.06", "0.133"),
    ("MET",  "1.07", "0.99\u20131.15", "0.077"),
    ("EGFR", "1.06", "0.97\u20131.17", "0.207"),
    ("SRC",  "1.07", "0.86\u20131.32", "0.556"),
    ("LCK",  "0.91", "0.80\u20131.03", "0.118"),
    ("YES1 *", "1.28", "1.03\u20131.60", "0.028"),
    ("FYN",  "0.86", "0.73\u20131.01", "0.063"),
    ("LYN",  "0.97", "0.80\u20131.17", "0.724"),
    ("DDR1", "0.84", "0.69\u20131.04", "0.104"),
    ("Age",  "1.01", "0.99\u20131.02", "0.429"),
]

table1 = doc.add_table(rows=len(t1_data), cols=4)
table1.style = 'Table Grid'
for r_idx, row_data in enumerate(t1_data):
    row = table1.rows[r_idx]
    for c_idx, cell_text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = cell_text
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text)
        if r_idx == 0:
            run.bold = True
        if r_idx == 6:  # YES1 row
            run.bold = True

doc.add_paragraph()
para(doc, "* YES1: significant at p < 0.05 in multivariate model.", indent=False)

page_break(doc)

# Table 2
h(doc, "Table 2. pARACNe Kinase\u2013Substrate Validation in CPTAC-LUAD (Top 10 Pairs, FDR\u00a0<\u00a00.05)", 2)
para(doc,
    "Kinase protein abundance (proteome) vs. substrate phosphorylation (phosphoproteome) Spearman "
    "correlation across 110 CPTAC-LUAD patients. All 752 tested triplets; FDR by "
    "Benjamini\u2013Hochberg.",
    indent=False)

t2_data = [
    ("Kinase", "Substrate", "Phosphosite", "Spearman \u03c1", "FDR"),
    ("EGFR",  "DBNL",     "S232",         "0.45",  "0.0006"),
    ("PTK2",  "MATR3",    "S604",         "\u22120.43", "0.0013"),
    ("EGFR",  "TJAP1",    "S300",         "0.41",  "0.0020"),
    ("EGFR",  "EPB41L1",  "S407",         "0.38",  "0.0078"),
    ("EGFR",  "BAIAP2L1", "S414",         "0.36",  "0.0090"),
    ("SRC",   "PRKCD",    "S525",         "0.36",  "0.0090"),
    ("YES1",  "COPA",     "S173",         "0.37",  "0.0090"),
    ("LYN",   "PYGL",     "S838",         "0.36",  "0.0090"),
    ("EGFR",  "DBNL",     "S282/S283",    "0.35",  "0.0136"),
]

table2 = doc.add_table(rows=len(t2_data), cols=5)
table2.style = 'Table Grid'
for r_idx, row_data in enumerate(t2_data):
    row = table2.rows[r_idx]
    for c_idx, cell_text in enumerate(row_data):
        cell = row.cells[c_idx]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(cell_text)
        if r_idx == 0:
            run.bold = True

page_break(doc)

# ── FIGURE LEGENDS ────────────────────────────────────────────────────────────
h(doc, "FIGURE LEGENDS", 1)

legends = [
    ("Figure 1. pARACNe Hub Kinase Identification.",
     "(A) Bar chart of substrate connectivity (number of unique substrates) for the top 20 "
     "tyrosine kinases in the pARACNe LUAD S2 protein-level network. Focal kinases "
     "(\u2265\u00a0100 substrates) are highlighted. (B) Substrate counts for the nine focal "
     "kinases with individual color coding. Network encompasses 46 kinases and 2,064 "
     "kinase\u2013substrate interactions."),

    ("Figure 2. Kinase Expression\u2013Drug Sensitivity Associations in GDSC2 LUAD Cell Lines.",
     "(A) Heatmap of Spearman \u03c1 values for all focal kinase \u00d7 top compound associations "
     "from full 286-drug GDSC2 scan (62 LUAD cell lines; 2,574 tests). Color scale: blue = "
     "negative correlation (high expression \u2192 sensitive), red = positive correlation "
     "(high expression \u2192 resistant). (B) Scatter plots of the top kinase\u2013drug pairs. "
     "Trend lines represent linear regression."),

    ("Figure 3. YES1 as an Independent Prognostic Factor in TCGA-LUAD.",
     "(A) Kaplan\u2013Meier overall survival curves for LCK and YES1 stratified by tertile "
     "expression (TCGA-LUAD, n\u00a0=\u00a0497; top vs. bottom 33%). Log-rank p-values shown. "
     "(B) Forest plot of Cox multivariate hazard ratios (n\u00a0=\u00a0487). Points represent HR; "
     "horizontal bars are 95% confidence intervals. YES1 is highlighted (p\u00a0=\u00a00.028)."),

    ("Figure 4. Composite Kinase Activation Score (KAS) in CPTAC-LUAD.",
     "(A) Violin plots of individual KAS for nine focal kinases in CPTAC-LUAD (n\u00a0=\u00a0110). "
     "KAS is defined as the mean log2 phosphorylation ratio of pARACNe-predicted substrate "
     "phosphosites. (B) PCA loading plot showing contribution of each kinase to PC1 of the "
     "multi-kinase KAS matrix (PC1\u00a0=\u00a080.8%). (C) PC1 vs PC2 scatter of 110 CPTAC "
     "patients. (D) KAS comparison by EGFR mutation status (Mann\u2013Whitney U test)."),

    ("Figure 5. SRC-Family Co-Activation and pARACNe Network Validation.",
     "(A) Kinase co-activation correlation matrix (Spearman \u03c1) computed from CPTAC "
     "phosphoproteomics. Note near-perfect co-activation among YES1, FYN, and LYN "
     "(\u03c1\u00a0>\u00a00.90). (B) Corresponding co-expression matrix from TCGA mRNA data. "
     "(C) Validation of pARACNe predictions: scatter plots of kinase protein abundance vs. "
     "substrate phosphorylation for top validated pairs. (D) Summary of validated "
     "kinase\u2013substrate interactions per kinase (FDR\u00a0<\u00a00.05 of 752 tested)."),
]

for title, body in legends:
    lp = doc.add_paragraph()
    lp.paragraph_format.first_line_indent = Cm(0)
    r = lp.add_run(title + " ")
    r.bold = True
    lp.add_run(body)
    doc.add_paragraph()

page_break(doc)

# ── FIGURES ───────────────────────────────────────────────────────────────────
h(doc, "FIGURES", 1)

PUB_DIR = os.path.join(RESULTS_DIR, "publication_figures")

figure_map = [
    ("Figure 1", os.path.join(PUB_DIR, "Figure1_HubKinases.png")),
    ("Figure 2", os.path.join(PUB_DIR, "Figure2_GDSC2_Correlations.png")),
    ("Figure 3", os.path.join(PUB_DIR, "Figure3_YES1_Survival.png")),
    ("Figure 4", os.path.join(PUB_DIR, "Figure4_KAS_CPTAC.png")),
    ("Figure 5", os.path.join(PUB_DIR, "Figure5_CoActivation_Validation.png")),
]

for fig_label, img_path in figure_map:
    fp = doc.add_paragraph()
    fp.paragraph_format.first_line_indent = Cm(0)
    r = fp.add_run(fig_label)
    r.bold = True; r.font.size = Pt(11)
    add_image_safe(doc, img_path, width_inches=6.2)
    page_break(doc)

# Supplementary Figure S1
h(doc, "SUPPLEMENTARY FIGURES", 1)
fp = doc.add_paragraph()
fp.paragraph_format.first_line_indent = Cm(0)
r = fp.add_run("Supplementary Figure S1")
r.bold = True; r.font.size = Pt(11)
add_image_safe(doc, os.path.join(PUB_DIR, "FigureS1_KM_Median_Grid.png"), width_inches=6.2)
page_break(doc)

# ── SUPPLEMENTARY TABLES (placeholders) ──────────────────────────────────────
h(doc, "SUPPLEMENTARY TABLES", 1)
for stitle in [
    "Table S1. Full GDSC2 kinase\u2013drug correlation results (2,574 tests). "
    "Available in LUAD_Faz2b_Results.xlsx.",
    "Table S2. GDSC2 ANOVA results for EGFR mutation vs. drug IC50 in LUAD. "
    "Available in LUAD_ANOVA_Sun Mar 1 09_17_41 2026.csv.",
    "Table S3. All pARACNe validated kinase\u2013substrate pairs (37 pairs, FDR\u00a0<\u00a00.05). "
    "Available in LUAD_Faz2_Results.xlsx.",
    "Table S4. TCGA-LUAD patient characteristics by YES1 tertile. To be generated.",
]:
    sp = doc.add_paragraph(style='Normal')
    sp.paragraph_format.first_line_indent = Cm(0)
    sp.add_run(stitle)

# ── SAVE ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
print(f"File size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
